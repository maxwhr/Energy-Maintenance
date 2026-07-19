from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


class DocumentParserError(ValueError):
    pass


@dataclass
class ParsedPage:
    page_number: int | None
    text: str


@dataclass
class ParsedDocument:
    text: str
    pages: list[ParsedPage] = field(default_factory=list)
    page_count: int | None = None
    metadata: dict[str, Any] = field(default_factory=dict)
    warnings: list[str] = field(default_factory=list)


class DocumentParser:
    def parse_document(self, file_path: str | Path, file_ext: str) -> ParsedDocument:
        path = Path(file_path)
        extension = file_ext.lower().lstrip(".")
        if extension in {"txt", "md"}:
            return self._parse_text(path, extension)
        if extension == "pdf":
            return self._parse_pdf(path)
        if extension == "docx":
            return self._parse_docx(path)
        raise DocumentParserError(f"Unsupported document extension: {extension}")

    def _parse_text(self, path: Path, extension: str) -> ParsedDocument:
        warnings: list[str] = []
        text: str | None = None
        for encoding in ("utf-8", "gbk"):
            try:
                text = path.read_text(encoding=encoding)
                if encoding != "utf-8":
                    warnings.append("File decoded with gbk fallback")
                break
            except UnicodeDecodeError:
                continue
        if text is None:
            raise DocumentParserError("Unable to decode text file with utf-8 or gbk")
        if not text.strip():
            raise DocumentParserError("Parsed document text is empty")
        return ParsedDocument(
            text=text,
            pages=[ParsedPage(page_number=None, text=text)],
            page_count=None,
            metadata={"parser": "plain_text", "extension": extension, "parser_version": "structured_parser_v1"},
            warnings=warnings,
        )

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise DocumentParserError("pypdf is not installed; PDF parsing is unavailable") from exc

        try:
            reader = PdfReader(str(path))
        except Exception as exc:
            raise DocumentParserError(f"Unable to read PDF file: {exc}") from exc

        pages: list[ParsedPage] = []
        warnings: list[str] = []
        for index, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text() or ""
            except Exception as exc:
                page_text = ""
                warnings.append(f"Page {index} parse failed: {exc}")
            if page_text.strip():
                pages.append(ParsedPage(page_number=index, text=page_text))

        text = "\n\n".join(page.text for page in pages)
        if not text.strip():
            raise DocumentParserError("Parsed PDF text is empty; scanned PDFs require OCR and are not supported")
        return ParsedDocument(
            text=text,
            pages=pages,
            page_count=len(reader.pages),
            metadata={"parser": "pypdf", "extension": "pdf", "parser_version": "structured_parser_v1"},
            warnings=warnings,
        )

    def _parse_docx(self, path: Path) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError as exc:
            raise DocumentParserError("python-docx is not installed; DOCX parsing is unavailable") from exc

        try:
            document = Document(str(path))
        except Exception as exc:
            raise DocumentParserError(f"Unable to read DOCX file: {exc}") from exc

        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                style = (paragraph.style.name or "").lower() if paragraph.style else ""
                if style.startswith("heading"):
                    try:
                        level = min(6, max(1, int(style.split()[-1])))
                    except (TypeError, ValueError):
                        level = 1
                    parts.append(f"{'#' * level} {text}")
                else:
                    parts.append(text)

        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    parts.append(" | ".join(values))

        image_count = len(document.inline_shapes)
        if image_count:
            parts.extend(f"[Image placeholder {index}]" for index in range(1, image_count + 1))
        text = "\n\n".join(parts)
        if not text.strip():
            raise DocumentParserError("Parsed DOCX text is empty")
        return ParsedDocument(
            text=text,
            pages=[ParsedPage(page_number=None, text=text)],
            page_count=None,
            metadata={"parser": "python-docx", "extension": "docx", "parser_version": "structured_parser_v1", "image_count": image_count},
            warnings=[],
        )
