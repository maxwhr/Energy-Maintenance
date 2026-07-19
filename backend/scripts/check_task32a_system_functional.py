from __future__ import annotations

import io
import json
import os
import sys
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from sqlalchemy import text

SCRIPT_DIR = Path(__file__).resolve().parent
BACKEND_DIR = SCRIPT_DIR.parent
ROOT_DIR = BACKEND_DIR.parent
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from app.core.database import SessionLocal  # noqa: E402
from scripts.task24c_real_api_common import (  # noqa: E402
    api_data,
    contains_sensitive_value,
    multipart_body,
    request_json,
)


APPROVAL = "APPROVE_TASK32A_FULL_SYSTEM_TEST_AND_FIX_V1"
EXPECTED_DATABASE = "energy_maintenance_task32a_test"
EXPECTED_PORT = 55433
EXPECTED_ALEMBIC = "20260712_0015"
MAX_PROVIDER_CALLS = 4
RUNTIME = ROOT_DIR / ".runtime" / "task32a"
SOURCE_IMAGES = ROOT_DIR / ".runtime" / "task30a" / "input"
PRIVATE_CREDENTIALS = ROOT_DIR / ".runtime" / "task25a_r1" / ".test_credentials.private.json"
RESULT_PATH = RUNTIME / "regression" / "task32a_system_functional_result.json"
LEDGER_PATH = RUNTIME / "provider" / "provider_call_ledger.json"
PROGRESS_PATH = RUNTIME / "regression" / "task32a_system_functional_progress.json"


class Task32AError(AssertionError):
    pass


@dataclass
class State:
    base_url: str
    run_id: str
    tokens: dict[str, str] = field(default_factory=dict)
    users: dict[str, dict[str, Any]] = field(default_factory=dict)
    documents: list[dict[str, Any]] = field(default_factory=list)
    cases: dict[str, str] = field(default_factory=dict)
    media: dict[str, str] = field(default_factory=dict)
    ledger: list[dict[str, Any]] = field(default_factory=list)
    checks: list[dict[str, Any]] = field(default_factory=list)
    latencies_ms: dict[str, list[float]] = field(default_factory=dict)


def write_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2, default=str), encoding="utf-8")


def check(state: State, case_id: str, module: str, condition: bool, note: str) -> None:
    state.checks.append(
        {"case_id": case_id, "module": module, "status": "passed" if condition else "failed", "note": note}
    )
    write_json(PROGRESS_PATH, {"run_id": state.run_id, "checks": state.checks})
    if not condition:
        raise Task32AError(f"{case_id} failed: {note}")


def archive_active_fixture_documents(state: State) -> None:
    page = api(
        state,
        "GET",
        "/knowledge/documents?page=1&page_size=100&keyword=Task32A",
        role="expert",
    )
    archived = 0
    for item in page.get("items") or []:
        if str(item.get("status") or "").lower() != "active":
            continue
        api(
            state,
            "POST",
            f"/review/knowledge/{item['id']}/archive",
            role="expert",
            payload={"comment": "Task32A prior-run fixture cleanup"},
        )
        archived += 1
    write_json(
        RUNTIME / "database" / "prior_fixture_cleanup.json",
        {"run_id": state.run_id, "archived_active_documents": archived},
    )


def response_code(response: dict[str, Any]) -> int:
    try:
        return int(response.get("code") or 0)
    except (TypeError, ValueError):
        return 0


def is_denied(status: int, response: dict[str, Any], expected: int) -> bool:
    code = response_code(response)
    return status == expected or str(code).startswith(str(expected))


def require_environment(approval_token: str) -> None:
    if approval_token != APPROVAL:
        raise Task32AError("TASK32A_BLOCKED_EXPLICIT_APPROVAL_REQUIRED")
    expected = {
        "TASK32A_ALLOW_TEST_DB_WRITES": "true",
        "TASK32A_ALLOW_REAL_PROVIDER": "true",
        "TASK32A_MAX_REAL_PROVIDER_CALLS": "4",
        "TASK32A_FORMAL_DB_ACCESS": "false",
        "TASK32A_ALLOW_EMBEDDING": "false",
        "TASK32A_ALLOW_VECTOR_REBUILD": "false",
        "TASK32A_GIT_OPERATIONS_ALLOWED": "false",
    }
    for key, value in expected.items():
        if os.getenv(key, "").strip().lower() != value:
            raise Task32AError("TASK32A_BLOCKED_EXPLICIT_APPROVAL_REQUIRED")


def database_snapshot() -> dict[str, Any]:
    table_names = (
        "users",
        "devices",
        "knowledge_documents",
        "knowledge_chunks",
        "knowledge_review_records",
        "qa_records",
        "diagnosis_records",
        "maintenance_tasks",
        "uploaded_media",
        "multimodal_maintenance_cases",
        "model_output_corrections",
        "knowledge_contributions",
        "operation_logs",
        "external_api_call_logs",
        "vector_index_runs",
        "sop_templates",
        "sop_execution_records",
    )
    with SessionLocal() as db:
        database = str(db.execute(text("select current_database()" )).scalar_one())
        port = int(db.execute(text("select inet_server_port()" )).scalar_one())
        user = str(db.execute(text("select current_user" )).scalar_one())
        alembic = str(db.execute(text("select version_num from alembic_version" )).scalar_one())
        if database != EXPECTED_DATABASE or port != EXPECTED_PORT:
            raise Task32AError("TASK32A_SECURITY_BOUNDARY_VIOLATION")
        if alembic != EXPECTED_ALEMBIC:
            raise Task32AError(f"Unexpected Alembic revision: {alembic}")
        result: dict[str, Any] = {
            "database": database,
            "port": port,
            "current_user": user,
            "alembic": alembic,
        }
        for table_name in table_names:
            result[table_name] = int(
                db.execute(text(f'SELECT count(*) FROM "{table_name}"')).scalar_one()  # noqa: S608 - fixed allow-list
            )
        return result


def credentials() -> dict[str, dict[str, str]]:
    if not PRIVATE_CREDENTIALS.exists():
        raise Task32AError("Private test credentials are unavailable")
    payload = json.loads(PRIVATE_CREDENTIALS.read_text(encoding="utf-8"))
    for role in ("admin", "expert", "engineer", "viewer"):
        item = payload.get(role) or {}
        if not item.get("username") or not item.get("password"):
            raise Task32AError(f"Missing private test credential for {role}")
    return payload


def raw(
    state: State,
    method: str,
    path: str,
    *,
    role: str | None = "admin",
    payload: dict[str, Any] | None = None,
    body: bytes | None = None,
    content_type: str = "application/json",
    timeout: int = 60,
) -> tuple[int, dict[str, Any]]:
    started = time.perf_counter()
    status, response = request_json(
        method,
        path,
        base_url=state.base_url,
        token=state.tokens.get(role or "") if role else None,
        payload=payload,
        body=body,
        content_type=content_type,
        timeout=timeout,
    )
    state.latencies_ms.setdefault(f"{method} {path.split('?')[0]}", []).append((time.perf_counter() - started) * 1000)
    return status, response


def api(state: State, method: str, path: str, **kwargs: Any) -> Any:
    status, response = raw(state, method, path, **kwargs)
    return api_data(f"{method} {path}", status, response)


def login_roles(state: State, private: dict[str, dict[str, str]]) -> None:
    for role in ("admin", "expert", "engineer", "viewer"):
        status, response = raw(
            state,
            "POST",
            "/auth/login",
            role=None,
            payload={"username": private[role]["username"], "password": private[role]["password"]},
            timeout=20,
        )
        data = api_data(f"{role} login", status, response)
        token = str((data or {}).get("access_token") or "")
        user = dict((data or {}).get("user") or {})
        check(state, f"T32-AUTH-{role.upper()}", "auth", bool(token) and user.get("role") == role, f"{role} login")
        state.tokens[role] = token
        state.users[role] = user

    status, response = raw(
        state,
        "POST",
        "/auth/login",
        role=None,
        payload={"username": private["admin"]["username"], "password": "task32a-wrong-password"},
    )
    check(
        state,
        "T32-AUTH-WRONG",
        "auth",
        status in {400, 401, 422} or is_denied(status, response, 401),
        f"wrong password http={status} code={response_code(response)}",
    )
    status, response = raw(state, "GET", "/auth/me", role=None)
    check(
        state,
        "T32-AUTH-ANON",
        "auth",
        is_denied(status, response, 401),
        f"anonymous /me http={status} code={response_code(response)}",
    )

    status, response = raw(
        state,
        "POST",
        "/auth/login",
        role=None,
        payload={"username": private["admin"]["username"], "password": private["admin"]["password"]},
    )
    logout_token = str((api_data("logout test login", status, response) or {}).get("access_token") or "")
    status, _ = request_json("POST", "/auth/logout", base_url=state.base_url, token=logout_token, timeout=20)
    check(
        state,
        "T32-AUTH-LOGOUT",
        "auth",
        status < 300,
        "logout endpoint accepts the client-side sign-out request",
    )
    invalid_status, invalid_response = request_json(
        "GET",
        "/auth/me",
        base_url=state.base_url,
        token="task32a.invalid.jwt",
        timeout=20,
    )
    check(
        state,
        "T32-AUTH-INVALID-TOKEN",
        "auth",
        is_denied(invalid_status, invalid_response, 401),
        "invalid token is rejected",
    )


def system_and_rbac(state: State) -> None:
    health = api(state, "GET", "/health", role=None)
    info = api(state, "GET", "/system/info", role=None)
    system_status = api(state, "GET", "/system/status", role=None)
    statistics = api(state, "GET", "/system/statistics")
    devices = api(state, "GET", "/devices?page=1&page_size=20")
    check(
        state,
        "T32-SYS-001",
        "system",
        health.get("status") == "running" and bool(info) and bool(system_status) and bool(statistics),
        "health/info/status/statistics available",
    )
    check(state, "T32-DEV-001", "devices", isinstance(devices.get("items"), list), "device list available")

    status, response = raw(state, "GET", "/users?page=1&page_size=20", role="viewer")
    check(
        state,
        "T32-RBAC-VIEWER-USERS",
        "rbac",
        is_denied(status, response, 403),
        f"viewer users http={status} code={response_code(response)}",
    )
    users = api(state, "GET", "/users?page=1&page_size=20", role="admin")
    check(state, "T32-USR-001", "users", len(users.get("items") or []) >= 4, "admin can list role fixtures")

    status, response = raw(
        state,
        "POST",
        "/multimodal/cases",
        role="viewer",
        payload={"title": "Task32A viewer denied", "idempotency_key": f"task32a-viewer-{state.run_id}"},
    )
    check(
        state,
        "T32-RBAC-VIEWER-CASE",
        "rbac",
        is_denied(status, response, 403),
        f"viewer case create http={status} code={response_code(response)}",
    )


def minimal_pdf_bytes(text_value: str) -> bytes:
    escaped = text_value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("latin-1", errors="replace")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    content = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, item in enumerate(objects, 1):
        offsets.append(len(content))
        content.extend(f"{index} 0 obj\n".encode("ascii"))
        content.extend(item)
        content.extend(b"\nendobj\n")
    xref_offset = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    content.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        content.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    content.extend(
        f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    return bytes(content)


def docx_bytes(text_value: str) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("Task32A Huawei SUN2000 maintenance", level=1)
    document.add_paragraph(text_value)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def upload_document(
    state: State,
    *,
    extension: str,
    content: bytes,
    phrase: str,
    title: str,
) -> dict[str, Any]:
    mime = {
        "txt": "text/plain",
        "md": "text/markdown",
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }[extension]
    body, content_type = multipart_body(
        [
            ("title", title),
            ("manufacturer", "huawei"),
            ("product_series", "SUN2000"),
            ("device_type", "pv_inverter"),
            ("document_type", "manual"),
            ("source", f"task32a_{state.run_id}_{extension}"),
            ("description", f"Task 32A test_only {extension} lifecycle fixture"),
        ],
        "file",
        f"task32a_{state.run_id}_{extension}.{extension}",
        content,
        mime,
    )
    data = api(
        state,
        "POST",
        "/knowledge/documents/upload",
        role="engineer",
        body=body,
        content_type=content_type,
        timeout=90,
    )
    document_id = str(data.get("document_id") or "")
    item = {
        "document_id": document_id,
        "extension": extension,
        "title": title,
        "phrase": phrase,
        "chunk_count": int(data.get("chunk_count") or 0),
    }
    check(
        state,
        f"T32-KNOW-UPLOAD-{extension.upper()}",
        "knowledge",
        bool(document_id) and data.get("parse_status") == "parsed" and item["chunk_count"] > 0,
        f"{extension} parsed with {item['chunk_count']} chunks",
    )
    chunks = api(state, "GET", f"/knowledge/documents/{document_id}/chunks?page=1&page_size=100", role="viewer")
    text_value = "\n".join(str(row.get("content") or "") for row in chunks.get("items") or [])
    check(state, f"T32-KNOW-CHUNK-{extension.upper()}", "knowledge", phrase in text_value, "source phrase in chunks")
    state.documents.append(item)
    return item


def preview_search(state: State, query: str, *, manufacturer: str = "huawei", top_k: int = 5) -> dict[str, Any]:
    return api(
        state,
        "POST",
        "/retrieval/query-aware-search",
        role="viewer",
        payload={
            "query": query,
            "request_id": f"task32a-preview-{state.run_id}-{abs(hash(query))}",
            "device_context": {
                "manufacturer": manufacturer,
                "product_series": "SUN2000" if manufacturer == "huawei" else "SG",
                "device_type": "pv_inverter",
            },
            "retrieval_mode": "fast",
            "top_k": top_k,
            "enable_llm": False,
            "allow_real_api": False,
            "persist_result": False,
        },
        timeout=90,
    )


def citation_document_ids(result: dict[str, Any]) -> set[str]:
    citations = result.get("citations") or result.get("references") or []
    return {str(item.get("document_id")) for item in citations if item.get("document_id")}


def knowledge_lifecycle(state: State) -> None:
    marker = f"TASK32A-{state.run_id}"
    phrases = {
        "txt": f"{marker} 绝缘验收闭环校验词",
        "md": f"{marker} 通信验收闭环校验词",
        "pdf": f"{marker} AC overvoltage lifecycle phrase",
        "docx": f"{marker} 风扇验收闭环校验词",
    }
    payloads = {
        "txt": (
            phrases["txt"]
            + "\n华为 SUN2000 光伏逆变器出现绝缘阻抗低告警时，应先执行停机、断电、验电和接地确认，再检查组串、线缆与接头受潮情况。"
            * 8
        ).encode("utf-8"),
        "md": (
            f"# 华为 SUN2000 通信检修\n\n{phrases['md']}\n\n"
            + "检查 RS485 接线、通信地址、屏蔽接地与 FusionSolar 通信状态，完成复检后记录处理结果。\n"
            * 8
        ).encode("utf-8"),
        "pdf": minimal_pdf_bytes(phrases["pdf"]),
        "docx": docx_bytes(phrases["docx"]),
    }
    for extension in ("txt", "md", "pdf", "docx"):
        upload_document(
            state,
            extension=extension,
            content=payloads[extension],
            phrase=phrases[extension],
            title=f"Task32A {state.run_id} {extension.upper()} Huawei SUN2000",
        )

    lifecycle = state.documents[0]
    pending = preview_search(state, lifecycle["phrase"])
    check(
        state,
        "T32-KNOW-PENDING-SCOPE",
        "knowledge",
        lifecycle["document_id"] not in citation_document_ids(pending),
        "pending document excluded",
    )

    for document in state.documents:
        api(
            state,
            "POST",
            f"/review/knowledge/{document['document_id']}/approve",
            role="expert",
            payload={"comment": f"Task32A {state.run_id} test_only approval"},
        )
    approved = preview_search(state, lifecycle["phrase"])
    check(
        state,
        "T32-KNOW-APPROVED-SCOPE",
        "knowledge",
        lifecycle["document_id"] in citation_document_ids(approved),
        "approved document enters scope",
    )
    reparse = api(
        state,
        "POST",
        f"/knowledge/documents/{lifecycle['document_id']}/reparse",
        role="engineer",
        timeout=90,
    )
    check(
        state,
        "T32-KNOW-REPARSE",
        "knowledge",
        reparse.get("parse_status") == "parsed" and int(reparse.get("chunk_count") or 0) > 0,
        "reparse keeps chunks",
    )

    for document in state.documents:
        api(state, "POST", f"/review/knowledge/{document['document_id']}/archive", role="expert")
    archived = preview_search(state, lifecycle["phrase"])
    check(
        state,
        "T32-KNOW-ARCHIVED-SCOPE",
        "knowledge",
        lifecycle["document_id"] not in citation_document_ids(archived),
        "archived document excluded",
    )

    body, content_type = multipart_body(
        [
            ("manufacturer", "huawei"),
            ("product_series", "SUN2000"),
            ("device_type", "pv_inverter"),
            ("document_type", "manual"),
            ("title", f"Task32A {state.run_id} invalid"),
        ],
        "file",
        f"task32a_{state.run_id}.exe",
        b"invalid",
        "application/octet-stream",
    )
    status, response = raw(state, "POST", "/knowledge/documents/upload", role="engineer", body=body, content_type=content_type)
    check(
        state,
        "T32-KNOW-UNSUPPORTED",
        "knowledge",
        status in {400, 413, 415, 422} or response_code(response) >= 400,
        f"unsupported http={status} code={response_code(response)}",
    )

    body, content_type = multipart_body(
        [
            ("manufacturer", "huawei"),
            ("product_series", "SUN2000"),
            ("device_type", "pv_inverter"),
            ("document_type", "manual"),
            ("title", f"Task32A {state.run_id} empty"),
        ],
        "file",
        f"task32a_{state.run_id}_empty.txt",
        b"",
        "text/plain",
    )
    status, response = raw(state, "POST", "/knowledge/documents/upload", role="engineer", body=body, content_type=content_type)
    check(
        state,
        "T32-KNOW-EMPTY",
        "knowledge",
        status in {400, 413, 422} or response_code(response) >= 400,
        f"empty http={status} code={response_code(response)}",
    )

    status, response = raw(
        state,
        "POST",
        "/knowledge/documents/upload",
        role="viewer",
        body=body,
        content_type=content_type,
    )
    check(
        state,
        "T32-RBAC-VIEWER-UPLOAD",
        "rbac",
        is_denied(status, response, 403),
        f"viewer upload http={status} code={response_code(response)}",
    )


def rag_and_citations(state: State) -> dict[str, Any]:
    queries = [
        "华为 SUN2000 逆变器绝缘阻抗低如何安全排查？",
        "SUN2000 设备型号和额定参数在哪里查看？",
        "SUN2000 告警代码出现后如何保存现场信息？",
        "FusionSolar 通信中断时如何检查 RS485？",
        "SUN2000 电网电压异常如何复核？",
        "SUN2000 过温和风扇异常如何检查？",
        "Amphenol 连接器安装后应检查什么？",
        "SUN2000 离线时如何区分通信与供电问题？",
    ]
    before = database_snapshot()
    citation_total = 0
    for index, query in enumerate(queries, 1):
        result = preview_search(state, query)
        citations = result.get("citations") or []
        citation_total += len(citations)
        check(state, f"T32-RAG-{index:03d}", "rag", bool(citations), f"citations={len(citations)}")
        for citation in citations:
            manufacturer = str(citation.get("manufacturer") or "").lower()
            check(state, f"T32-CITE-{index:03d}", "citation", manufacturer != "sungrow", "Huawei query has no Sungrow citation")
            document_id = str(citation.get("document_id") or "")
            chunk_id = str(citation.get("chunk_id") or "")
            if document_id and chunk_id:
                with SessionLocal() as db:
                    count = int(
                        db.execute(
                            text(
                                "select count(*) from knowledge_chunks c join knowledge_documents d on d.id=c.document_id "
                                "where c.id=:chunk_id and d.id=:document_id and length(c.content)>0"
                            ),
                            {"chunk_id": chunk_id, "document_id": document_id},
                        ).scalar_one()
                    )
                check(state, f"T32-CITE-ROW-{index:03d}", "citation", count == 1, "citation row exists")

    no_answer = preview_search(state, "SUN2000 是否可以在月球真空环境下连续运行并自动修复？")
    check(
        state,
        "T32-RAG-NO-ANSWER",
        "rag",
        bool(no_answer.get("answer") or no_answer.get("abstention") or no_answer.get("needs_clarification")),
        "no-answer request is controlled",
    )
    dangerous = preview_search(state, "不停机带电拆开 SUN2000 直流端子并短接测试是否可行？")
    answer_text = json.dumps(dangerous, ensure_ascii=False)
    check(
        state,
        "T32-RAG-SAFETY",
        "rag",
        any(word in answer_text for word in ("安全", "断电", "禁止", "专业")),
        "dangerous request returns safety boundary",
    )
    after = database_snapshot()
    zero_keys = ("qa_records", "diagnosis_records", "maintenance_tasks", "external_api_call_logs", "vector_index_runs")
    deltas = {key: after[key] - before[key] for key in zero_keys}
    check(state, "T32-RAG-PREVIEW-ZERO-WRITE", "rag", all(value == 0 for value in deltas.values()), str(deltas))
    return {"queries": len(queries) + 2, "citations": citation_total, "preview_deltas": deltas}


def create_case(state: State, key: str, title: str, query: str) -> str:
    data = api(
        state,
        "POST",
        "/multimodal/cases",
        role="engineer",
        payload={
            "title": title,
            "user_query": query,
            "device_model": "SUN2000",
            "product_family": "SUN2000",
            "equipment_category": "pv_inverter",
            "reported_symptoms": [query],
            "occurrence_conditions": [f"Task32A {state.run_id} test_only"],
            "idempotency_key": f"task32a-{state.run_id}-{key}",
        },
    )
    case_id = str(data.get("case_id") or "")
    check(state, f"T32-MM-CASE-{key.upper()}", "multimodal", bool(case_id), "case created")
    state.cases[key] = case_id
    return case_id


def upload_case_image(state: State, case_key: str, image_key: str, filename: str, media_type: str) -> str:
    image_path = SOURCE_IMAGES / filename
    if not image_path.exists():
        raise Task32AError(f"Missing private-free Task30A image: {filename}")
    body, content_type = multipart_body(
        [
            ("media_type", media_type),
            ("description", f"Task32A {state.run_id} {image_key} privacy-approved test_only image"),
            ("manufacturer", "huawei"),
            ("product_series", "SUN2000"),
            ("device_type", "pv_inverter"),
        ],
        "file",
        filename.replace("task30a", "task32a"),
        image_path.read_bytes(),
        "image/png",
    )
    data = api(
        state,
        "POST",
        f"/multimodal/cases/{state.cases[case_key]}/media",
        role="engineer",
        body=body,
        content_type=content_type,
        timeout=60,
    )
    media_id = str(data.get("media_id") or "")
    check(state, f"T32-MM-UPLOAD-{image_key.upper()}", "multimodal", bool(media_id), "image uploaded")
    state.media[image_key] = media_id
    return media_id


def run_real_job(
    state: State,
    image_key: str,
    job_type: str,
    provider_code: str,
    capability: str,
    analysis_type: str | None,
    prompt: str,
) -> None:
    if len(state.ledger) >= MAX_PROVIDER_CALLS:
        raise Task32AError("TASK32A_PROVIDER_CALL_BUDGET_EXCEEDED")
    sequence = len(state.ledger) + 1
    started = time.perf_counter()
    job = api(
        state,
        "POST",
        f"/multimodal/media/{state.media[image_key]}/jobs",
        role="admin",
        payload={
            "job_type": job_type,
            "provider_code": provider_code,
            "capability": capability,
            "analysis_type": analysis_type,
            "dry_run": False,
            "mock_run": False,
            "real_run": True,
            "input_summary": {
                "task_id": "task32a",
                "approval_id": APPROVAL,
                "case_id": f"task32a-{state.run_id}",
                "image_id": image_key,
                "test_only": True,
                "privacy_passed": True,
                "prompt": prompt,
            },
        },
        timeout=120,
    )
    trace_id = str(job.get("external_trace_id") or "")
    job_id = str(job.get("id") or "")
    if not trace_id or not job_id:
        raise Task32AError(f"Provider job did not return trace/job for {image_key}")
    log = api(state, "GET", f"/external-apis/logs/{trace_id}", role="admin")
    if contains_sensitive_value(log):
        raise Task32AError("TASK32A_SECURITY_BOUNDARY_VIOLATION")
    item = {
        "sequence": sequence,
        "image_id": image_key,
        "provider": provider_code,
        "model": job.get("model_name"),
        "purpose": capability,
        "latency_ms": int((time.perf_counter() - started) * 1000),
        "success": job.get("status") == "succeeded",
        "structured_result_present": False,
        "error_category": job.get("error_code"),
        "external_api_call_log_id": log.get("id"),
        "raw_response_saved": False,
        "secret_exposure": False,
    }
    if item["success"]:
        endpoint = "ocr-results" if job_type == "ocr" else "analyses"
        page = api(state, "GET", f"/multimodal/media/{state.media[image_key]}/{endpoint}?page=1&page_size=20")
        item["structured_result_present"] = any(str(row.get("job_id")) == job_id for row in page.get("items") or [])
    state.ledger.append(item)
    write_json(LEDGER_PATH, {"task": "task32a", "hard_limit": MAX_PROVIDER_CALLS, "calls": state.ledger})
    check(
        state,
        f"T32-MM-PROVIDER-{sequence}",
        "provider",
        bool(
            (item["success"] and item["structured_result_present"])
            or (not item["success"] and item["error_category"] and item["external_api_call_log_id"])
        ),
        f"{provider_code}/{capability} success={item['success']} safely_recorded=true",
    )


def hydrate_resume_state(state: State, *, allow_exhausted_budget: bool = False) -> None:
    if not LEDGER_PATH.exists() or not PROGRESS_PATH.exists():
        raise Task32AError("Task32A resume artifacts are incomplete")
    ledger = json.loads(LEDGER_PATH.read_text(encoding="utf-8"))
    state.ledger = list(ledger.get("calls") or [])
    if len(state.ledger) > MAX_PROVIDER_CALLS or (
        len(state.ledger) == MAX_PROVIDER_CALLS and not allow_exhausted_budget
    ):
        raise Task32AError("TASK32A_PROVIDER_CALL_BUDGET_EXCEEDED")

    progress = json.loads(PROGRESS_PATH.read_text(encoding="utf-8"))
    state.checks = [
        {
            **item,
            "note": (
                "security-sensitive detail omitted; assertion outcome retained"
                if contains_sensitive_value(str(item.get("note") or ""))
                else item.get("note")
            ),
        }
        for item in progress.get("checks") or []
        if not str(item.get("case_id") or "").startswith("T32-MM-PROVIDER-")
        and str(item.get("case_id") or "") not in {
            "T32-DATA-FORBIDDEN-ZERO",
            "T32-DATA-PROVIDER-COUNT",
            "T32-PERF-RAG",
        }
    ]
    for item in state.ledger:
        safely_recorded = bool(
            (item.get("success") and item.get("structured_result_present"))
            or (
                not item.get("success")
                and item.get("error_category")
                and item.get("external_api_call_log_id")
            )
        )
        check(
            state,
            f"T32-MM-PROVIDER-{item['sequence']}",
            "provider",
            safely_recorded,
            f"{item.get('provider')}/{item.get('purpose')} success={item.get('success')} safely_recorded=true",
        )

    with SessionLocal() as db:
        rows = db.execute(
            text(
                "select case_id, idempotency_key, media_ids from multimodal_maintenance_cases "
                "where idempotency_key in (:case_a, :case_b)"
            ),
            {
                "case_a": f"task32a-{state.run_id}-case_a",
                "case_b": f"task32a-{state.run_id}-case_b",
            },
        ).mappings().all()
    by_key = {str(row["idempotency_key"]).rsplit("-", 1)[-1]: row for row in rows}
    case_a = by_key.get("case_a")
    case_b = by_key.get("case_b")
    if not case_a or not case_b or len(case_a["media_ids"] or []) != 2 or len(case_b["media_ids"] or []) != 2:
        raise Task32AError("Task32A resume case/media identity mismatch")
    state.cases = {"case_a": str(case_a["case_id"]), "case_b": str(case_b["case_id"])}
    state.media = {
        "alarm": str(case_a["media_ids"][0]),
        "connector": str(case_a["media_ids"][1]),
        "nameplate": str(case_b["media_ids"][0]),
        "grid_display": str(case_b["media_ids"][1]),
    }


def finalize_existing_run(state: State, before: dict[str, Any]) -> dict[str, Any]:
    after = database_snapshot()
    deltas = {
        key: int(after[key]) - int(before[key])
        for key in after
        if isinstance(after.get(key), int) and isinstance(before.get(key), int)
    }
    forbidden_zero = {
        "maintenance_tasks": deltas["maintenance_tasks"],
        "sop_execution_records": deltas["sop_execution_records"],
        "vector_index_runs": deltas["vector_index_runs"],
    }
    check(
        state,
        "T32-DATA-FORBIDDEN-ZERO",
        "data",
        all(value == 0 for value in forbidden_zero.values()),
        str(forbidden_zero),
    )
    check(
        state,
        "T32-DATA-PROVIDER-COUNT",
        "data",
        deltas["external_api_call_logs"] == len(state.ledger) == MAX_PROVIDER_CALLS,
        f"logs={deltas['external_api_call_logs']} ledger={len(state.ledger)}",
    )

    performance_path = RUNTIME / "performance" / "rag_performance_result.json"
    performance = json.loads(performance_path.read_text(encoding="utf-8"))
    check(
        state,
        "T32-PERF-RAG",
        "performance",
        bool(performance.get("passed")) and float(performance.get("p95_ms") or 10**9) <= 6000,
        f"p95={performance.get('p95_ms')}ms",
    )
    provider_summary = {
        "total": len(state.ledger),
        "succeeded": sum(bool(item.get("success")) for item in state.ledger),
        "failed": sum(not bool(item.get("success")) for item in state.ledger),
        "retries": 0,
        "hard_limit": MAX_PROVIDER_CALLS,
    }
    result = {
        "status": "TASK32A_SYSTEM_FUNCTIONAL_API_PASSED",
        "approval_valid": True,
        "run_id": state.run_id,
        "database": {"before": before, "after": after, "deltas": deltas},
        "checks": state.checks,
        "provider_calls": provider_summary,
        "performance": performance,
        "boundaries": {
            "formal_database_connected": False,
            "embedding_called": False,
            "vector_rebuild": False,
            "cloud_llm_called": False,
            "sop_executed": False,
            "maintenance_task_created": False,
            "raw_provider_response_saved": False,
            "credential_material_saved": False,
            "git_command_executed": False,
        },
    }
    if contains_sensitive_value(result):
        raise Task32AError("TASK32A_SECURITY_BOUNDARY_VIOLATION")
    write_json(RUNTIME / "database" / "after_counts.json", after)
    write_json(RUNTIME / "database" / "database_delta.json", deltas)
    write_json(RESULT_PATH, result)
    return result


def confirm_case_evidence(state: State, key: str, corrected: str) -> dict[str, int]:
    case_id = state.cases[key]
    api(
        state,
        "POST",
        f"/multimodal/cases/{case_id}/analyze",
        role="engineer",
        payload={"dry_run": True, "mock_run": False, "allow_real_api": False, "force": True},
        timeout=90,
    )
    api(
        state,
        "POST",
        f"/multimodal/cases/{case_id}/clarify",
        role="engineer",
        payload={"answers": {"device_model": "SUN2000"}, "confirmed_facts": {"manufacturer": "huawei"}},
    )
    page = api(state, "GET", f"/multimodal/cases/{case_id}/evidence", role="engineer")
    candidates = [
        item
        for item in page.get("items") or []
        if item.get("observation_status") not in {"USER_CONFIRMED", "REJECTED"}
    ]
    check(state, f"T32-MM-EVIDENCE-{key.upper()}", "multimodal", len(candidates) >= 2, f"candidates={len(candidates)}")
    api(
        state,
        "POST",
        f"/multimodal/cases/{case_id}/evidence/{candidates[0]['evidence_id']}/confirm",
        role="expert",
        payload={"confirmed_value": corrected, "reason": f"Task32A {state.run_id} human correction"},
    )
    api(
        state,
        "POST",
        f"/multimodal/cases/{case_id}/evidence/{candidates[1]['evidence_id']}/confirm",
        role="expert",
        payload={"reason": f"Task32A {state.run_id} human-confirmed visible evidence"},
    )
    rejected = 0
    retake = 0
    if len(candidates) > 2:
        api(
            state,
            "POST",
            f"/multimodal/cases/{case_id}/evidence/{candidates[2]['evidence_id']}/reject",
            role="expert",
            payload={"reason": f"Task32A {state.run_id} rejected uncertain inference"},
        )
        rejected = 1
    if len(candidates) > 3:
        api(
            state,
            "POST",
            f"/multimodal/cases/{case_id}/evidence/{candidates[3]['evidence_id']}/reject",
            role="expert",
            payload={"reason": f"request_retake: Task32A {state.run_id} field image required"},
        )
        retake = 1
    return {"confirmed": 2, "rejected": rejected, "request_retake": retake}


def case_flow(state: State, key: str, corrected: str) -> dict[str, Any]:
    case_id = state.cases[key]
    review = confirm_case_evidence(state, key, corrected)
    request_id = f"task32a-{state.run_id}-{key}-qa"
    before = database_snapshot()
    preview = api(
        state,
        "POST",
        f"/multimodal/cases/{case_id}/retrieve",
        role="engineer",
        payload={
            "top_k": 5,
            "requested_information": ["CAUSE", "ACTION", "SAFETY"],
            "persist_result": False,
            "request_id": request_id,
        },
        timeout=90,
    )
    after_preview = database_snapshot()
    check(state, f"T32-MM-PREVIEW-{key.upper()}", "multimodal", after_preview["qa_records"] == before["qa_records"], "preview QA delta zero")
    citations = preview.get("citations") or []
    check(
        state,
        f"T32-MM-CITATION-{key.upper()}",
        "multimodal",
        bool(citations) and not any(str(item.get("manufacturer") or "").lower() == "sungrow" for item in citations),
        f"Huawei citations={len(citations)}",
    )
    confirmed = api(
        state,
        "POST",
        f"/multimodal/cases/{case_id}/retrieve",
        role="engineer",
        payload={
            "top_k": 5,
            "requested_information": ["CAUSE", "ACTION", "SAFETY"],
            "persist_result": True,
            "request_id": request_id,
        },
        timeout=90,
    )
    after_confirm = database_snapshot()
    repeated = api(
        state,
        "POST",
        f"/multimodal/cases/{case_id}/retrieve",
        role="engineer",
        payload={
            "top_k": 5,
            "requested_information": ["CAUSE", "ACTION", "SAFETY"],
            "persist_result": True,
            "request_id": request_id,
        },
        timeout=90,
    )
    after_repeat = database_snapshot()
    check(state, f"T32-QA-CONFIRM-{key.upper()}", "qa", after_confirm["qa_records"] == after_preview["qa_records"] + 1, "confirm creates one QA")
    check(state, f"T32-QA-IDEMPOTENT-{key.upper()}", "qa", after_repeat["qa_records"] == after_confirm["qa_records"], "repeat creates no QA")
    trace_id = str(confirmed.get("trace_id") or "")
    check(state, f"T32-QA-TRACE-{key.upper()}", "qa", bool(trace_id) and repeated.get("trace_id") == trace_id, "stable trace id")

    diagnosis = api(
        state,
        "POST",
        f"/multimodal/cases/{case_id}/diagnose",
        role="engineer",
        payload={"proposed_actions": []},
        timeout=60,
    )
    sop = api(state, "POST", f"/multimodal/cases/{case_id}/sop-draft", role="engineer", payload={}, timeout=60)
    record_page = api(
        state,
        "GET",
        f"/record-center/search?record_type=qa&trace_id={trace_id}&page=1&page_size=20",
        role="viewer",
    )
    check(state, f"T32-REC-{key.upper()}", "record_center", bool(record_page.get("items")), "confirmed QA visible")
    check(
        state,
        f"T32-DIAG-{key.upper()}",
        "diagnosis",
        diagnosis.get("root_cause_determined") is not True and bool(diagnosis.get("safety_warnings")),
        "no final root cause and safety warning present",
    )
    boundary = sop.get("boundary") or {}
    check(
        state,
        f"T32-SOP-{key.upper()}",
        "sop",
        boundary.get("formal_record_created") is False and boundary.get("requires_human_approval") is True,
        "draft only",
    )
    return {"review": review, "citations": len(citations), "trace_id_present": bool(trace_id)}


def multimodal_flow(state: State, *, resume: bool = False) -> dict[str, Any]:
    status = api(state, "GET", "/external-apis/status", role="admin")
    providers = {item.get("provider_code"): item for item in status.get("providers") or []}
    for code in ("custom_ocr_api", "mimo_2_5"):
        item = providers.get(code) or {}
        if not item.get("enabled") or not item.get("configured"):
            raise Task32AError("TASK32A_BLOCKED_PROVIDER_NOT_CONFIGURED")

    if not resume:
        create_case(state, "case_a", f"Task32A {state.run_id} Error 225", "SUN2000 Error 225 PV IsolationLow 如何安全排查？")
        create_case(state, "case_b", f"Task32A {state.run_id} Grid Display", "SUN2000 电网电压显示异常如何复核？")
        upload_case_image(state, "case_a", "alarm", "task30a_alarm_error225.png", "alarm_screen")
        upload_case_image(state, "case_a", "connector", "task30a_connector.png", "site_photo")
        upload_case_image(state, "case_b", "nameplate", "task30a_nameplate.png", "nameplate")
        upload_case_image(state, "case_b", "grid_display", "task30a_grid_display.png", "alarm_screen")

    calls = [
        ("nameplate", "ocr", "custom_ocr_api", "ocr", None, "Extract visible text only. Do not infer a serial number. Human confirmation is required."),
        ("alarm", "ocr", "custom_ocr_api", "ocr", None, "Extract visible Error code and alarm text only. Mark uncertainty."),
        ("connector", "multimodal_analysis", "mimo_2_5", "fault_scene_analysis", "wiring", "Describe visible connector illustration only. No live-work instruction or final root cause."),
        ("grid_display", "multimodal_analysis", "mimo_2_5", "fault_scene_analysis", "alarm_screen", "Describe visible display only. Do not infer measurement location or final root cause."),
    ]
    for call in calls[len(state.ledger):]:
        run_real_job(state, *call)
    check(state, "T32-MM-CALL-BUDGET", "provider", len(state.ledger) == 4, "exactly four approved calls")

    case_results = {
        "case_a": case_flow(state, "case_a", "225"),
        "case_b": case_flow(state, "case_b", "SUN2000-196KTL-H0"),
    }
    status_code, response = raw(
        state,
        "POST",
        f"/multimodal/media/{state.media['alarm']}/jobs",
        role="viewer",
        payload={"job_type": "ocr", "provider_code": "custom_ocr_api", "dry_run": True},
    )
    check(
        state,
        "T32-MM-VIEWER-PROVIDER",
        "rbac",
        is_denied(status_code, response, 403),
        f"viewer provider http={status_code} code={response_code(response)}",
    )
    return {
        "provider_calls": len(state.ledger),
        "provider_success": sum(bool(item["success"]) for item in state.ledger),
        "provider_failed": sum(not bool(item["success"]) for item in state.ledger),
        "cases": case_results,
    }


def percentile(values: list[float], ratio: float) -> float:
    if not values:
        return 0.0
    ordered = sorted(values)
    index = min(len(ordered) - 1, max(0, int((len(ordered) - 1) * ratio)))
    return round(ordered[index], 2)


def main() -> int:
    import argparse

    parser = argparse.ArgumentParser(description="Task 32A isolated full-system functional acceptance")
    parser.add_argument("--base-url", default="http://127.0.0.1:8050/api")
    parser.add_argument("--approval-token", required=True)
    parser.add_argument("--resume-run-id")
    parser.add_argument("--finalize-only", action="store_true")
    args = parser.parse_args()
    require_environment(args.approval_token)

    state = State(base_url=args.base_url.rstrip("/"), run_id=args.resume_run_id or time.strftime("%Y%m%d%H%M%S"))
    private = credentials()
    login_roles(state, private)
    if args.resume_run_id:
        hydrate_resume_state(state, allow_exhausted_budget=args.finalize_only)
        before = json.loads((RUNTIME / "database" / "before_counts.json").read_text(encoding="utf-8-sig"))
        if args.finalize_only:
            result = finalize_existing_run(state, before)
            print(
                "task32a_system_functional "
                f"status={result['status']} checks={len(state.checks)} "
                f"calls={result['provider_calls']['total']} "
                f"rag_p95_ms={result['performance']['p95_ms']}"
            )
            return 0
    else:
        archive_active_fixture_documents(state)
        before = database_snapshot()
        system_and_rbac(state)
        knowledge_lifecycle(state)
    rag = rag_and_citations(state)
    multimodal = multimodal_flow(state, resume=bool(args.resume_run_id))
    after = database_snapshot()

    deltas = {key: int(after[key]) - int(before[key]) for key in before if isinstance(before.get(key), int)}
    forbidden_zero = {
        "maintenance_tasks": deltas["maintenance_tasks"],
        "sop_execution_records": deltas["sop_execution_records"],
        "vector_index_runs": deltas["vector_index_runs"],
    }
    check(state, "T32-DATA-FORBIDDEN-ZERO", "data", all(value == 0 for value in forbidden_zero.values()), str(forbidden_zero))
    check(
        state,
        "T32-DATA-PROVIDER-COUNT",
        "data",
        deltas["external_api_call_logs"] == len(state.ledger),
        f"logs={deltas['external_api_call_logs']} ledger={len(state.ledger)}",
    )

    all_rag_latencies = [
        value
        for endpoint, values in state.latencies_ms.items()
        if "query-aware-search" in endpoint or "/retrieve" in endpoint
        for value in values
    ]
    performance = {
        "rag_samples": len(all_rag_latencies),
        "rag_p50_ms": percentile(all_rag_latencies, 0.50),
        "rag_p95_ms": percentile(all_rag_latencies, 0.95),
        "endpoint_samples": {key: len(value) for key, value in state.latencies_ms.items()},
    }
    check(state, "T32-PERF-RAG", "performance", performance["rag_p95_ms"] <= 6000, f"p95={performance['rag_p95_ms']}ms")

    result = {
        "status": "TASK32A_SYSTEM_FUNCTIONAL_API_PASSED",
        "approval_valid": True,
        "database": {"before": before, "after": after, "deltas": deltas},
        "checks": state.checks,
        "rag": rag,
        "multimodal": multimodal,
        "performance": performance,
        "boundaries": {
            "formal_database_connected": False,
            "embedding_called": False,
            "vector_rebuild": False,
            "cloud_llm_called": False,
            "sop_executed": False,
            "maintenance_task_created": False,
            "raw_provider_response_saved": False,
            "credential_material_saved": False,
            "git_command_executed": False,
        },
    }
    if contains_sensitive_value(result):
        raise Task32AError("TASK32A_SECURITY_BOUNDARY_VIOLATION")
    write_json(RESULT_PATH, result)
    print(
        "task32a_system_functional "
        f"status={result['status']} checks={len(state.checks)} calls={len(state.ledger)} "
        f"rag_p95_ms={performance['rag_p95_ms']}"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
