from __future__ import annotations

import io
import json
import os
import re
import socket
import sys
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any
from urllib import request

from sqlalchemy import create_engine, text

SCRIPT_DIR = Path(__file__).resolve().parent
BACKEND_DIR = SCRIPT_DIR.parent
ROOT_DIR = BACKEND_DIR.parent

if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

import check_global_acceptance as cga  # noqa: E402
import check_real_frontend_api_integration as task21a  # noqa: E402


API_BASE_URL = os.getenv("TASK21B_API_BASE_URL", "http://127.0.0.1:8010/api").rstrip("/")
APP_BASE_URL = API_BASE_URL[:-4] if API_BASE_URL.endswith("/api") else API_BASE_URL
DATABASE_URL = os.getenv("DATABASE_URL", "")
RUN_ID = os.getenv("TASK21B_RUN_ID", time.strftime("%Y%m%d%H%M%S"))
MARKER = f"Task21B_{RUN_ID}"
RESULT_PATH = ROOT_DIR / ".runtime" / "21B_remaining_real_test_result.json"

cga.API_BASE_URL = API_BASE_URL
cga.APP_BASE_URL = APP_BASE_URL
cga.ADMIN_USERNAME = os.getenv("TASK21B_ADMIN_USERNAME") or os.getenv("FULL_SMOKE_ADMIN_USERNAME", "admin")
cga.ADMIN_PASSWORD = os.getenv("TASK21B_ADMIN_PASSWORD") or os.getenv("FULL_SMOKE_ADMIN_PASSWORD", "admin123456")
cga.TEST_PASSWORD = os.getenv("TASK21B_TEST_PASSWORD", "Task21B_pass123")
cga.RUN_ID = RUN_ID
cga.MARKER = MARKER
task21a.API_BASE_URL = API_BASE_URL
task21a.APP_BASE_URL = APP_BASE_URL


class Task21BError(AssertionError):
    pass


@dataclass(slots=True)
class TestDocument:
    document_id: str
    title: str
    ext: str
    phrase: str
    chunk_count: int
    trace_id: str | None = None


@dataclass(slots=True)
class Task21BState:
    admin_token: str | None = None
    expert_token: str | None = None
    engineer_token: str | None = None
    viewer_token: str | None = None
    users: dict[str, dict[str, Any]] = field(default_factory=dict)
    documents: list[TestDocument] = field(default_factory=list)
    pending_document_id: str | None = None
    pending_trace_id: str | None = None
    approved_trace_id: str | None = None
    cleanup_notes: list[str] = field(default_factory=list)
    frontend_button_matrix: list[dict[str, Any]] = field(default_factory=list)
    port_status: dict[str, Any] = field(default_factory=dict)
    db_status: dict[str, Any] = field(default_factory=dict)


def port_is_open(port: int, host: str = "127.0.0.1", timeout: float = 1.5) -> bool:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
        sock.settimeout(timeout)
        return sock.connect_ex((host, port)) == 0


def safe_get_json(url: str) -> dict[str, Any] | None:
    try:
        with request.urlopen(url, timeout=5) as response:
            return json.loads(response.read().decode("utf-8"))
    except Exception:  # noqa: BLE001
        return None


def check_environment(state: Task21BState, recorder: cga.Recorder) -> None:
    status: dict[str, Any] = {}
    for port in (8000, 8010, 5432, 55432):
        status[str(port)] = {"open": port_is_open(port)}

    openapi_8000 = safe_get_json("http://127.0.0.1:8000/openapi.json") if status["8000"]["open"] else None
    openapi_api = safe_get_json(f"{APP_BASE_URL}/openapi.json")
    health = safe_get_json(f"{API_BASE_URL}/health")
    status["8000"]["openapi_title"] = (openapi_8000 or {}).get("info", {}).get("title")
    status["api_base"] = API_BASE_URL
    status["app_base"] = APP_BASE_URL
    status["app_openapi_title"] = (openapi_api or {}).get("info", {}).get("title")
    status["health_name"] = ((health or {}).get("data") or {}).get("name")
    state.port_status = status

    if status["health_name"] != "Energy-Maintenance":
        raise Task21BError(f"{API_BASE_URL}/health is not Energy-Maintenance: {status['health_name']!r}")
    if "Energy-Maintenance" not in str(status["app_openapi_title"]):
        raise Task21BError(f"{APP_BASE_URL}/openapi.json is not Energy-Maintenance: {status['app_openapi_title']!r}")
    recorder.add(
        "environment",
        "ports and backend identity",
        "passed",
        f"8000={status['8000']['openapi_title']!r}, 8010={status['app_openapi_title']!r}, 5432_open={status['5432']['open']}, 55432_open={status['55432']['open']}",
    )


def check_database(state: Task21BState, recorder: cga.Recorder) -> None:
    if not DATABASE_URL:
        raise Task21BError("DATABASE_URL is not set for direct PostgreSQL verification")
    engine = create_engine(DATABASE_URL, pool_pre_ping=True)
    with engine.connect() as connection:
        row = connection.execute(
            text(
                "select current_database() as db_name, current_user as db_user, inet_server_port() as db_port, 1 as ok"
            )
        ).mappings().one()
    state.db_status = dict(row)
    if state.db_status.get("db_name") != "energy_maintenance":
        raise Task21BError(f"unexpected database: {state.db_status}")
    recorder.add(
        "environment",
        "direct PostgreSQL connection",
        "passed",
        f"database={state.db_status['db_name']}, user={state.db_status['db_user']}, port={state.db_status['db_port']}",
    )


def ensure_task21b_users(state: Task21BState, recorder: cga.Recorder) -> None:
    admin_token, admin_user = cga.login(cga.ADMIN_USERNAME, cga.ADMIN_PASSWORD)
    state.admin_token = admin_token
    state.users["admin"] = admin_user
    recorder.add("auth", "admin login", "passed", f"user={admin_user.get('username')}")

    for role in ("engineer", "expert", "viewer"):
        username = f"{MARKER}_{role}"
        user = cga.ensure_user(admin_token, username, role, f"{MARKER} {role}")
        state.users[role] = user
        role_token, _ = cga.login(username, cga.TEST_PASSWORD)
        setattr(state, f"{role}_token", role_token)
        recorder.add("auth", f"{role} ready", "passed", f"id={user.get('id')}")


def minimal_pdf_bytes(text_value: str) -> bytes:
    escaped = text_value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("latin-1", errors="replace")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    content = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for idx, obj in enumerate(objects, 1):
        offsets.append(len(content))
        content.extend(f"{idx} 0 obj\n".encode("ascii"))
        content.extend(obj)
        content.extend(b"\nendobj\n")
    xref_offset = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    content.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        content.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    content.extend(
        f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    return bytes(content)


def docx_bytes(text_value: str) -> bytes:
    try:
        from docx import Document
    except Exception as exc:  # noqa: BLE001
        raise Task21BError(f"python-docx is unavailable for DOCX upload test: {exc}") from exc
    document = Document()
    document.add_heading("Task21B SUN2000 maintenance note", level=1)
    document.add_paragraph(text_value)
    document.add_paragraph("Check DC insulation, cabinet grounding, and alarm reset conditions.")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def file_payloads() -> dict[str, tuple[str, bytes, str, str]]:
    phrases = {
        "txt": f"{MARKER} TXT SUN2000 insulation alarm phrase",
        "md": f"{MARKER} MD Sungrow SG grid alarm phrase",
        "docx": f"{MARKER} DOCX Huawei FusionSolar fan inspection phrase",
        "pdf": f"{MARKER} PDF SUN2000 AC overvoltage phrase",
    }
    return {
        "txt": (
            f"{MARKER}_sun2000.txt",
            (phrases["txt"] + "\nPV inverter maintenance TXT parse verification.").encode("utf-8"),
            "text/plain",
            phrases["txt"],
        ),
        "md": (
            f"{MARKER}_sungrow.md",
            (f"# Sungrow SG maintenance\n\n{phrases['md']}\n\nCheck grid voltage and relay status.").encode("utf-8"),
            "text/markdown",
            phrases["md"],
        ),
        "docx": (
            f"{MARKER}_fusionsolar.docx",
            docx_bytes(phrases["docx"]),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            phrases["docx"],
        ),
        "pdf": (
            f"{MARKER}_sun2000.pdf",
            minimal_pdf_bytes(phrases["pdf"]),
            "application/pdf",
            phrases["pdf"],
        ),
    }


def upload_document(
    state: Task21BState,
    *,
    ext: str,
    filename: str,
    content: bytes,
    mime: str,
    phrase: str,
    manufacturer: str,
    product_series: str,
    title_suffix: str,
) -> TestDocument:
    fields = [
        ("title", f"{MARKER} {title_suffix} {ext.upper()}"),
        ("manufacturer", manufacturer),
        ("product_series", product_series),
        ("device_type", "pv_inverter"),
        ("document_type", "manual"),
        ("source", f"{MARKER}_{ext}_source"),
        ("description", f"Task21B {ext} upload verification."),
    ]
    body, content_type = cga.multipart_body(fields, "file", filename, content, mime)
    status, response = cga.request_json(
        "POST",
        "/knowledge/documents/upload",
        token=state.engineer_token,
        body=body,
        content_type=content_type,
        timeout=90,
    )
    data = cga.require_dict(cga.api_data(f"upload {ext}", status, response), f"upload {ext}")
    document_id = str(data.get("document_id") or "")
    chunk_count = int(data.get("chunk_count") or 0)
    if not document_id or data.get("parse_status") != "parsed" or chunk_count <= 0:
        raise Task21BError(f"{ext} upload did not parse: {data}")
    return TestDocument(
        document_id=document_id,
        title=str(fields[0][1]),
        ext=ext,
        phrase=phrase,
        chunk_count=chunk_count,
    )


def verify_document_chunks(state: Task21BState, doc: TestDocument) -> None:
    status, response = cga.request_json(
        "GET",
        f"/knowledge/documents/{doc.document_id}/chunks?page=1&page_size=20",
        token=state.viewer_token,
    )
    page = cga.require_page(cga.api_data(f"{doc.ext} chunks", status, response), f"{doc.ext} chunks")
    combined = "\n".join(str(item.get("content", "")) for item in page["items"])
    if doc.phrase not in combined:
        raise Task21BError(f"{doc.ext} chunks did not contain expected phrase")


def approve_document(state: Task21BState, doc: TestDocument) -> None:
    status, response = cga.request_json(
        "POST",
        f"/review/knowledge/{doc.document_id}/approve",
        token=state.expert_token,
        payload={"comment": f"Task21B approve {doc.ext}."},
    )
    data = cga.require_dict(cga.api_data(f"approve {doc.ext}", status, response), f"approve {doc.ext}")
    document = data.get("document") if isinstance(data.get("document"), dict) else data
    if document.get("review_status") != "approved":
        raise Task21BError(f"{doc.ext} document was not approved: {data}")


def retrieval_for_phrase(
    state: Task21BState,
    phrase: str,
    *,
    manufacturer: str | None = None,
    product_series: str | None = None,
    document_type: str = "manual",
) -> dict[str, Any]:
    payload = {
        "query": phrase,
        "manufacturer": manufacturer,
        "product_series": product_series,
        "device_type": "pv_inverter",
        "document_type": document_type,
        "top_k": 5,
        "include_sources": True,
        "enable_model_enhancement": False,
    }
    status, response = cga.request_json("POST", "/retrieval/query", token=state.viewer_token, payload=payload, timeout=90)
    return cga.require_dict(cga.api_data("retrieval query", status, response), "retrieval query")


def check_multi_format_uploads(state: Task21BState, recorder: cga.Recorder) -> None:
    payloads = file_payloads()
    mapping = {
        "txt": ("huawei", "SUN2000", "Huawei SUN2000"),
        "md": ("sungrow", "SG", "Sungrow SG"),
        "docx": ("huawei", "FusionSolar", "Huawei FusionSolar"),
        "pdf": ("huawei", "SUN2000", "Huawei SUN2000"),
    }
    for ext, (filename, content, mime, phrase) in payloads.items():
        manufacturer, product_series, title_suffix = mapping[ext]
        doc = upload_document(
            state,
            ext=ext,
            filename=filename,
            content=content,
            mime=mime,
            phrase=phrase,
            manufacturer=manufacturer,
            product_series=product_series,
            title_suffix=title_suffix,
        )
        verify_document_chunks(state, doc)
        approve_document(state, doc)
        result = retrieval_for_phrase(
            state,
            phrase,
            manufacturer=manufacturer,
            product_series=product_series,
        )
        references = result.get("references") or []
        if not any(str(ref.get("document_id")) == doc.document_id for ref in references):
            raise Task21BError(f"{ext} retrieval did not reference uploaded document")
        doc.trace_id = str(result.get("trace_id") or "")
        state.documents.append(doc)
        recorder.add(
            "knowledge multi-format",
            f"{ext} upload parse retrieval",
            "passed",
            f"document_id={doc.document_id}, chunks={doc.chunk_count}, trace_id={doc.trace_id}",
        )


def check_review_retrieval_effect(state: Task21BState, recorder: cga.Recorder) -> None:
    pending_phrase = f"{MARKER} pending review should not appear in retrieval"
    pending_doc = upload_document(
        state,
        ext="txt",
        filename=f"{MARKER}_pending.txt",
        content=(pending_phrase + "\nPending review document must not be retrieved before approval.").encode("utf-8"),
        mime="text/plain",
        phrase=pending_phrase,
        manufacturer="huawei",
        product_series="SUN2000",
        title_suffix="Pending Review",
    )
    state.pending_document_id = pending_doc.document_id

    before = retrieval_for_phrase(state, pending_phrase, manufacturer="huawei", product_series="SUN2000")
    before_refs = before.get("references") or []
    if any(str(ref.get("document_id")) == pending_doc.document_id for ref in before_refs):
        raise Task21BError("pending_review document appeared in retrieval before approval")
    recorder.add("knowledge review", "pending document excluded from retrieval", "passed", f"trace_id={before.get('trace_id')}")

    approve_document(state, pending_doc)
    after = retrieval_for_phrase(state, pending_phrase, manufacturer="huawei", product_series="SUN2000")
    after_refs = after.get("references") or []
    if not any(str(ref.get("document_id")) == pending_doc.document_id for ref in after_refs):
        raise Task21BError("approved document did not appear in retrieval")
    state.approved_trace_id = str(after.get("trace_id") or "")
    recorder.add("knowledge review", "approved document included in retrieval", "passed", f"trace_id={state.approved_trace_id}")

    status, response = cga.request_json("POST", f"/knowledge/documents/{pending_doc.document_id}/reparse", token=state.engineer_token)
    reparse = cga.require_dict(cga.api_data("pending document reparse", status, response), "pending document reparse")
    if reparse.get("parse_status") != "parsed" or int(reparse.get("chunk_count") or 0) <= 0:
        raise Task21BError(f"reparse failed for approved document: {reparse}")
    recorder.add("knowledge reparse/archive", "reparse approved document", "passed", f"chunk_count={reparse.get('chunk_count')}")

    status, response = cga.request_json("DELETE", f"/knowledge/documents/{pending_doc.document_id}", token=state.engineer_token)
    archived = cga.require_dict(cga.api_data("archive pending document", status, response), "archive pending document")
    if archived.get("status") != "archived":
        raise Task21BError(f"archive did not return archived status: {archived}")
    state.cleanup_notes.append(f"archived review-effect document: {pending_doc.document_id}")
    state.pending_document_id = None

    after_archive = retrieval_for_phrase(state, pending_phrase, manufacturer="huawei", product_series="SUN2000")
    archive_refs = after_archive.get("references") or []
    if any(str(ref.get("document_id")) == pending_doc.document_id for ref in archive_refs):
        raise Task21BError("archived document still appeared in retrieval")
    recorder.add("knowledge reparse/archive", "archived document excluded from retrieval", "passed", f"trace_id={after_archive.get('trace_id')}")


def check_permissions(state: Task21BState, recorder: cga.Recorder) -> None:
    fields = [
        ("title", f"{MARKER} viewer forbidden upload"),
        ("manufacturer", "huawei"),
        ("product_series", "SUN2000"),
        ("device_type", "pv_inverter"),
        ("document_type", "manual"),
    ]
    body, content_type = cga.multipart_body(
        fields,
        "file",
        f"{MARKER}_viewer_forbidden.txt",
        b"viewer upload should be forbidden",
        "text/plain",
    )
    status, response = cga.request_json(
        "POST",
        "/knowledge/documents/upload",
        token=state.viewer_token,
        body=body,
        content_type=content_type,
        timeout=30,
    )
    cga.expect_forbidden("viewer upload document", status, response)
    recorder.add("permissions", "viewer upload blocked", "passed")

    if state.documents:
        status, response = cga.request_json(
            "POST",
            f"/review/knowledge/{state.documents[0].document_id}/archive",
            token=state.viewer_token,
            payload={"comment": "viewer should not archive"},
        )
        cga.expect_forbidden("viewer archive review document", status, response)
        recorder.add("permissions", "viewer archive review blocked", "passed")

    status, response = cga.request_json(
        "POST",
        "/maintenance/tasks",
        token=state.viewer_token,
        payload={
            "title": f"{MARKER} viewer task forbidden",
            "fault_type": "low_insulation_resistance",
            "priority": "low",
            "fault_description": "viewer should not create task",
        },
    )
    cga.expect_forbidden("viewer create maintenance task", status, response)
    recorder.add("permissions", "viewer task create blocked", "passed")


def scan_frontend_buttons(state: Task21BState, recorder: cga.Recorder) -> None:
    scan = task21a.scan_frontend_api_calls(task21a.fetch_openapi_paths())
    api_functions = {call.function for call in scan.frontend_api_calls if call.function != "<inline>"}
    matrix: list[dict[str, Any]] = []
    for view_path in sorted((ROOT_DIR / "frontend" / "src" / "views").rglob("*.vue")):
        text_value = view_path.read_text(encoding="utf-8", errors="replace")
        click_handlers = re.findall(r"@click(?:\.stop)?(?:\.prevent)?=\"([^\"]+)\"", text_value)
        submit_handlers = re.findall(r"@submit(?:\.prevent)?=\"([^\"]+)\"", text_value)
        router_links = len(re.findall(r"<RouterLink\b", text_value))
        api_used = sorted(name for name in api_functions if re.search(rf"\b{name}\s*\(", text_value))
        matrix.append(
            {
                "page": str(view_path.relative_to(ROOT_DIR)).replace("\\", "/"),
                "click_handlers": len(click_handlers),
                "submit_handlers": len(submit_handlers),
                "router_links": router_links,
                "api_functions": api_used,
            }
        )
    state.frontend_button_matrix = matrix
    pages_with_actions = [item for item in matrix if item["click_handlers"] or item["submit_handlers"] or item["router_links"]]
    pages_with_api = [item for item in matrix if item["api_functions"]]
    if scan.missing_calls:
        raise Task21BError("frontend API scan has missing backend matches")
    recorder.add(
        "frontend alternative button check",
        "page actions mapped to real API functions",
        "passed",
        f"pages_with_actions={len(pages_with_actions)}, pages_with_api={len(pages_with_api)}, api_calls={len(scan.frontend_api_calls)}",
    )


def cleanup(state: Task21BState, recorder: cga.Recorder) -> None:
    archived_ids: set[str] = set()
    for doc in state.documents:
        status, response = cga.request_json("DELETE", f"/knowledge/documents/{doc.document_id}", token=state.engineer_token)
        if status < 400 and response.get("code") in (0, 200):
            archived_ids.add(doc.document_id)
            state.cleanup_notes.append(f"archived {doc.ext} document: {doc.document_id}")
        else:
            recorder.add("cleanup", f"archive {doc.ext} document", "partial", f"http={status}, message={response.get('message')}")
    if state.pending_document_id and state.pending_document_id not in archived_ids:
        status, response = cga.request_json(
            "DELETE",
            f"/knowledge/documents/{state.pending_document_id}",
            token=state.engineer_token,
        )
        if status < 400 and response.get("code") in (0, 200):
            state.cleanup_notes.append(f"archived pending/review-effect document: {state.pending_document_id}")
        else:
            recorder.add(
                "cleanup",
                "archive pending/review-effect document",
                "partial",
                f"http={status}, message={response.get('message')}",
            )
    for role, user in state.users.items():
        if role == "admin":
            continue
        user_id = user.get("id")
        if not user_id:
            continue
        status, response = cga.request_json("POST", f"/users/{user_id}/disable", token=state.admin_token)
        if status < 400 and response.get("code") in (0, 200):
            state.cleanup_notes.append(f"disabled user: {user.get('username')}")
        else:
            recorder.add("cleanup", f"disable {role}", "partial", f"http={status}, message={response.get('message')}")
    recorder.add("cleanup", "soft cleanup", "passed", "; ".join(state.cleanup_notes) or "no cleanup actions")


def run() -> tuple[cga.Recorder, Task21BState]:
    recorder = cga.Recorder()
    state = Task21BState()
    recorder.step("environment", "port and backend check", lambda: check_environment(state, recorder) or "")
    recorder.step("environment", "database direct check", lambda: check_database(state, recorder) or "")
    recorder.step("auth", "real role setup", lambda: ensure_task21b_users(state, recorder) or "")
    recorder.step("frontend alternative button check", "source action matrix", lambda: scan_frontend_buttons(state, recorder) or "")
    recorder.step("knowledge multi-format", "txt/md/docx/pdf upload parse retrieval", lambda: check_multi_format_uploads(state, recorder) or "")
    recorder.step("knowledge review", "review status retrieval effect", lambda: check_review_retrieval_effect(state, recorder) or "")
    recorder.step("permissions", "viewer boundary probes", lambda: check_permissions(state, recorder) or "")
    recorder.step("cleanup", "archive test data and disable users", lambda: cleanup(state, recorder) or "")
    return recorder, state


def summary_payload(recorder: cga.Recorder, state: Task21BState) -> dict[str, Any]:
    return {
        "status": "failed" if recorder.failed() else "passed",
        "api_base_url": API_BASE_URL,
        "app_base_url": APP_BASE_URL,
        "database_url_used": re.sub(r"(postgresql\+psycopg://[^:]+:)[^@]+(@.*)", r"\1***\2", DATABASE_URL),
        "run_id": RUN_ID,
        "marker": MARKER,
        "summary": {
            "total": len(recorder.results),
            "passed": recorder.count("passed"),
            "blocked": recorder.count("blocked"),
            "partial": recorder.count("partial"),
            "failed": recorder.count("failed"),
            "skipped": recorder.count("skipped"),
        },
        "environment": {
            "ports": state.port_status,
            "database": state.db_status,
        },
        "documents": [
            {
                "document_id": doc.document_id,
                "ext": doc.ext,
                "title": doc.title,
                "chunk_count": doc.chunk_count,
                "trace_id": doc.trace_id,
            }
            for doc in state.documents
        ],
        "review_status_effect": {
            "document_id": state.pending_document_id,
            "approved_trace_id": state.approved_trace_id,
        },
        "frontend_button_matrix": state.frontend_button_matrix,
        "cleanup_notes": state.cleanup_notes,
        "results": [
            {"area": result.area, "name": result.name, "status": result.status, "notes": result.notes}
            for result in recorder.results
        ],
    }


def main() -> int:
    try:
        recorder, state = run()
    except Exception as exc:  # noqa: BLE001
        recorder = cga.Recorder()
        state = Task21BState()
        recorder.add("Task21B", "fatal setup error", "failed", str(exc))
    payload = summary_payload(recorder, state)
    RESULT_PATH.parent.mkdir(parents=True, exist_ok=True)
    RESULT_PATH.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(payload, ensure_ascii=False, indent=2))
    return 1 if recorder.failed() else 0


if __name__ == "__main__":
    raise SystemExit(main())
